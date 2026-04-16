"""Generate CP650 study MCQs Word document."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Title ---
t = doc.add_heading("CP650 Interaction Design — Study Guide & MCQ Bank", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Based on Rogers, Sharp & Preece (6th ed.) lecture slides (Lectures 1–15).")
doc.add_paragraph("Correct answers are listed in the Answer Key at the end.")
doc.add_paragraph()

doc.add_heading("Part A — Lecture summaries (quick review)", level=1)
summaries = [
    ("Lecture 1", "Interaction design definition, goals, UX vs usability, multidisciplinary fields (6 academic disciplines, 5 interdisciplinary fields, 5 design practices), usability goals (6), team work, business examples."),
    ("Lecture 2", "Accessibility, inclusiveness, disability types (sensory, physical, cognitive), impairment (permanent, temporary, situational), design principles: visibility, feedback, constraints, mapping, consistency (aesthetic/internal/external), affordances, perceived affordances, simplicity."),
    ("Lecture 3", "ID process, Double Diamond, problem space, user involvement, user-centered approach (cognitive/behavioral/anthropomorphic/attitudinal), four core activities, Agile UX, patterns, tools, documentation."),
    ("Lecture 4", "Proof of concept, assumptions vs claims, conceptual models, metaphors, five interaction types, interface styles, paradigms/visions/theories/models/frameworks, Norman three-part model (designer, system image, user)."),
    ("Lecture 5", "Cognition: experiential/reflective, fast/slow thinking, attention, perception, memory, learning, language; Tullis, driving/phone, recognition vs recall, Miller 7±2 critique, PIM, MFA, design implications."),
    ("Lecture 6", "Social interaction, Sacks turn-taking (3 rules), remote conversation, VideoWindow, telepresence, social/co-presence, awareness (peripheral, situational), ReflectTable, engagement, chatbot ethics example."),
    ("Lecture 7", "Emotional UX, Norman three levels (visceral, behavioral, reflective), expressive interfaces, affective computing, persuasive tech, anthropomorphism."),
    ("Lecture 8", "Interfaces — styles and dialogue patterns (large topic)."),
    ("Lectures 9–10", "Data gathering: five key issues, interviews (unstructured/structured/semi-structured), focus groups (3–10), observation, questionnaires, ethics, pilots, triangulation."),
    ("Lectures 11–12", "Qualitative analysis, grounded theory, frameworks, presenting findings."),
    ("Lecture 13", "Discovering requirements."),
    ("Lecture 14", "Data at scale."),
    ("Lecture 15", "Ideate, design, prototyping, construction."),
]
for title, body in summaries:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)
doc.add_page_break()

doc.add_heading("Part B — Multiple choice questions", level=1)
doc.add_paragraph("Choose the best answer for each question.")

# (question_text, options [A,B,C,D], correct_index 0-3)
MCQS = [
    # Counts & lists
    ("How many academic disciplines are listed as contributing to interaction design (Psychology, Social Sciences, etc.)?", ["Four", "Five", "Six", "Seven"], 2),
    ("How many interdisciplinary fields are listed (HCI, Ubiquitous Computing, Human Factors, etc.)?", ["Three", "Four", "Five", "Six"], 2),
    ("How many design practices are listed (graphic, product, artist-design, industrial, film)?", ["Four", "Five", "Six", "Seven"], 1),
    ("How many usability goals are listed in Lecture 1?", ["Four", "Five", "Six", "Seven"], 2),
    ("In Hassenzahl’s UX model (slides), how many main types are named (pragmatic and hedonic)?", ["One", "Two", "Three", "Four"], 1),
    ("How many main disability types are listed (sensory, physical, cognitive)?", ["Two", "Three", "Four", "Five"], 1),
    ("How many impairment time categories are given (permanent, temporary, situational)?", ["Two", "Three", "Four", "Five"], 1),
    ("How many cognitive process areas are listed under “Cognitive processes” in Lecture 5?", ["Four", "Five", "Six", "Seven"], 1),
    ("How many basic turn-taking rules does Sacks et al. propose?", ["Two", "Three", "Four", "Five"], 1),
    ("How many interaction types are defined in Lecture 4 (instructing through responding)?", ["Four", "Five", "Six", "Seven"], 1),
    ("In Norman’s model (designer / system / user), how many interacting components?", ["Two", "Three", "Four", "Five"], 1),
    ("How many types of consistency are named (aesthetic, internal, external)?", ["Two", "Three", "Four", "Five"], 1),
    ("How many key issues for planning data gathering are in Lecture 9?", ["Three", "Four", "Five", "Six"], 2),
    ("Typical focus group size in the slides?", ["1–2 people", "3–10 people", "15–20 people", "25–30 people"], 1),
    ("Norman’s emotional design levels in Lecture 7?", ["Two", "Three", "Four", "Five"], 1),
    ("In the user-centered approach slide, how many kinds of user characteristics are named for early study?", ["Two", "Three", "Four", "Five"], 2),
    ("Lecture 5 lists two main types of learning:", ["Reflective and experiential", "Fast and slow", "Incidental and intentional", "Automatic and conscious"], 2),
    ("Characteristics of interaction design in Lecture 1 include how many main bullet themes (users involved, goals documented, iteration)?", ["Two", "Three", "Four", "Five"], 1),
    ("Pragmatic and hedonic dimensions appear in which author’s user experience model cited in the slides?", ["Kahneman", "Miller", "Hassenzahl", "Sacks"], 2),
    # Definitions L1
    ("Interaction design (Sharp, Rogers, Preece) is best described as:", ["Only visual design for websites", "Designing interactive products to support how people communicate and interact in everyday and working lives", "Only programming user interfaces", "Only marketing research"], 1),
    ("Winograd (1997) describes interaction design as:", ["Design of algorithms", "Design of spaces for human communication and interaction", "Design of databases", "Design of hardware only"], 1),
    ("Usability in the slides means:", ["Only speed of tasks", "Easy to learn, effective to use, and provides an enjoyable experience", "Only aesthetic beauty", "Only low cost"], 1),
    ("Goals of interaction design include:", ["Avoid users entirely", "Develop usable products and involve users in the design process", "Maximize lines of code", "Design only for experts"], 1),
    ("According to the slides, user experience:", ["Can be fully designed and guaranteed by the designer", "Cannot be fully “designed”—designers design for a user experience", "Is unrelated to emotion", "Applies only to software, not physical products"], 1),
    ("Garrett (2010) is cited to say:", ["Only apps have UX", "Every product that someone uses has a user experience", "UX applies only to websites", "UX is the same as usability"], 1),
    ("Nielsen and Norman (2014) define UX as relating to:", ["Only button colors", "All aspects of the end-user’s interaction with the company, services, and products", "Only loading time", "Only accessibility audits"], 1),
    ("The textbook cited on the slides is:", ["Norman (1988) only", "Rogers, Sharp, and Preece — 6th edition", "Kahneman (2011) only", "Sacks et al. (1978) only"], 1),
    ("Pragmatic quality in Hassenzahl’s model refers to:", ["How flashy the animation is", "How simple and practical it is for users to achieve their goals", "How many social media shares", "Server response time only"], 1),
    ("Hedonic quality in Hassenzahl’s model refers to:", ["Database normalization", "How stimulating the interaction is to users", "Keyboard layout only", "Legal compliance"], 1),
    # L1 examples & concepts
    ("The marble answering machine (Bishop, 1995) illustrates good design because it:", ["Requires many hidden steps", "Is based on how everyday objects behave and uses one-step core actions", "Has no sound", "Uses only voice commands"], 1),
    ("The elevator bad-design example: why are mistakes more likely on the bottom row?", ["Bottom buttons are larger", "Labels and controls look the same, so labels can be pushed by mistake", "Top row is disabled", "Bottom row uses voice only"], 1),
    ("The bad vending machine example stresses:", ["Bills are too small", "User must push a button before inserting money, unlike usual “insert bill first” habit", "It is too colorful", "It has no screen"], 1),
    ("The TiVo remote is praised for:", ["Having 100 identical buttons", "Peanut shape, logical layout, color-coded distinctive buttons", "No tactile feedback", "Requiring a phone app"], 1),
    ("Multidisciplinary teams in ID:", ["Always agree instantly", "Bring more ideas but can be hard to communicate across backgrounds", "Should exclude designers", "Never include engineers"], 1),
    ("Nielsen Norman Group is described as helping companies:", ["Avoid users", "Design human-centered products and services", "Write kernel code", "Sell only hardware"], 1),
    ("IDEO is cited as:", ["A database company", "Creating products, services, and environments for companies pioneering new value", "A law firm", "A CPU manufacturer"], 1),
    ("Cooper is cited as focusing on:", ["Random UI colors", "Goal-related design from research to product", "Only print design", "Only game physics"], 1),
    # Usability goals list
    ("Which is NOT listed as a usability goal in Lecture 1?", ["Effective to use", "Efficient to use", "Maximum advertising clicks", "Safe to use"], 2),
    ("“Easy to remember how to use” is:", ["A user experience goal only", "A usability goal", "Not mentioned", "Same as hedonic"], 1),
    # L2 accessibility & principles
    ("Accessibility refers to:", ["How fast the network is", "The extent to which a product is accessible by as many people as possible", "Only screen resolution", "Only paid features"], 1),
    ("Inclusiveness means:", ["Designing for one persona only", "Accommodating the widest possible number of people", "Excluding older adults", "English-only interfaces"], 1),
    ("Technologies can be inaccessible because they:", ["Always use large fonts", "Assume a type of interaction impossible for some impairments", "Never use sound", "Are always open source"], 1),
    ("Visibility as a principle means users can:", ["Never see advanced options", "See the state of the device and possible actions", "Only hear audio", "Skip all feedback"], 1),
    ("Feedback includes:", ["Only printed manuals", "Sound, highlighting, animation, and combinations", "Only error codes", "Nothing visual"], 1),
    ("Constraints help by:", ["Adding random button order", "Restricting possible actions to prevent incorrect selections", "Removing all labels", "Hiding the power button"], 1),
    ("Good mapping (e.g., color-coded keyboard/mouse ports) helps users:", ["Forget which cable goes where", "Associate controls with the right connectors logically", "Disable USB", "Remove audio"], 1),
    ("Internal consistency means:", ["Different behavior in every dialog", "Operations behave the same within an application", "Only external branding", "Random shortcuts"], 1),
    ("External consistency means sameness:", ["Only inside one dialog box", "Across applications and devices (often hard to achieve)", "Only for icons, never wording", "Only on Tuesdays"], 1),
    ("The phone vs calculator keypad layout example illustrates:", ["Perfect global consistency", "A case of external inconsistency", "Accessibility law", "Affective computing"], 1),
    ("Affordance (Norman) refers to:", ["Legal liability", "Attributes that suggest how to use an object", "Wi-Fi strength", "File compression"], 1),
    ("For virtual interfaces, Norman argues we should emphasize:", ["Physical affordances identical to wood and metal", "Perceived affordances and learned conventions", "No learnability", "Random placement"], 1),
    ("Jakob Nielsen’s simplicity idea (1999) suggests:", ["Add every possible widget", "Remove elements that are not needed if the design still works—balance with aesthetics", "Never use images", "Use only monospace fonts"], 1),
    # L3 process
    ("Interaction design as a process is focused on:", ["Ignoring requirements", "Discovering requirements, designing, prototyping, evaluating—focused on users and trade-offs", "Only final polish", "One-shot delivery"], 1),
    ("Approaches mentioned include:", ["Only waterfall with no users", "User-centered, activity-centered, and systems design", "Only graphic design", "Only sales-led design"], 1),
    ("Exploring the problem space asks:", ["Only about server cost", "Current UX, why change is needed, and how change improves the situation", "Only competitor logos", "Only font choice"], 1),
    ("Expectation management includes:", ["Overpromising features", "Realistic expectations, timely training, communication without hype", "No training", "Hiding release dates"], 1),
    ("Ownership in user involvement means:", ["Users own the company stock", "Users as active stakeholders—more acceptance if problems arise", "No user contact", "Only post-launch surveys"], 1),
    ("User-centered approach includes empirical measurement, meaning:", ["Guessing preferences", "Observing and analyzing users with scenarios, manuals, simulations, prototypes", "Only A/B without ethics", "Only expert inspection"], 1),
    ("Iterative design means:", ["Ship once and never fix", "When problems appear in testing, fix and test again", "Delete all prototypes", "Skip evaluation"], 1),
    ("Agile development timeboxes are described as about:", ["5–10 years", "One to three week iterations", "Exactly one hour", "No schedule"], 1),
    ("UX / technical debt arises from:", ["Too much user testing", "Short-term compromises and neglected UX that cause long-term problems", "Using prototypes", "Inclusive design"], 1),
    ("Agile UX and interaction design share:", ["No iteration", "Iteration, user involvement, measurable criteria", "No user involvement", "Only documentation"], 1),
    ("A parallel tracks approach in Agile UX suggests:", ["Design must finish after all coding", "Do UX work about one iteration ahead of development (or similar)", "No upfront vision", "Ban user research"], 1),
    ("Design patterns in software/OOD are:", ["Random screenshots", "Named problem/solution descriptions guiding design with trade-offs", "Only color palettes", "Only legal templates"], 1),
    ("Using design patterns everywhere without thought can lead to:", ["Always ideal code", "Over-architecture or unbalanced, inefficient code", "Fewer abstractions", "No reuse"], 1),
    # L4 conceptualizing
    ("A proof of concept helps teams:", ["Skip feasibility", "Scrutinize assumptions—feasibility, desirability, realism", "Avoid stakeholders", "Hide risks"], 1),
    ("An assumption is:", ["A proven theorem", "Something taken for granted that needs investigation", "A user persona", "A finished design"], 1),
    ("A claim is:", ["Always proven", "Stated as true but still open to question", "Same as a metaphor", "Only a bug report"], 1),
    ("A conceptual model is:", ["Source code", "A high-level description of how a system is organized and operates", "Only a logo", "Only a marketing slogan"], 1),
    ("Interface metaphors help users by:", ["Forcing memorization of commands", "Leveraging familiar knowledge to understand unfamiliar functionality", "Removing all visuals", "Banning touch"], 1),
    ("Problems with metaphors can include:", ["They always match culture", "They can break cultural rules, over-constrain imagination, or copy bad designs", "They never help novices", "They remove task concepts"], 1),
    ("The five interaction types are:", ["Type, click, swipe, pinch, zoom", "Instructing, conversing, manipulating, exploring, responding", "Red, green, blue, alpha, depth", "Waterfall, spiral, agile, lean, devops"], 1),
    ("“Instructing” fits tasks such as:", ["Only VR sculpting", "Repetitive actions like file management or spell-check", "Only passive alerts", "Only exploration worlds"], 1),
    ("“Conversing” fits:", ["Only spreadsheets", "Systems where interaction is like dialogue—chatbots, voice menus, help", "Only batch jobs", "Only hardware CAD"], 1),
    ("Direct manipulation benefits include:", ["Users never see results", "Novices learn quickly; immediate feedback; users feel in control", "Always requires long manuals", "No benefit for intermittent users"], 1),
    ("A disadvantage of direct manipulation can be:", ["It is always fastest for every task", "Some tasks are better delegated; can consume screen space; pointer can be slower than keys for experts", "It never works for novices", "It bans graphics"], 1),
    ("“Exploring” involves:", ["Only typing SQL", "Moving through virtual or physical environments and datasets", "Only printing", "Only system-initiated alerts"], 1),
    ("“Responding” (Lueg et al., 2018) means:", ["User always starts every action", "The system initiates interaction based on context, presence, or learned behavior", "Only email", "Only command line"], 1),
    ("A theory in HCI (Lecture 4) can address factors such as:", ["Only network latency", "Cognitive, social, and affective aspects", "Only typography", "Only stock markets"], 1),
    ("Examples of newer paradigms listed include:", ["Only mainframes", "Ubiquitous, pervasive, wearable computing, IoT", "Only punch cards", "Only fax machines"], 1),
    ("The designer’s model in Norman’s framework is:", ["What marketing posts on Twitter", "The designer’s idea of how the system should work", "The user’s mental model only", "The compiler output"], 1),
    ("The system image is:", ["Hidden source code only", "What is communicated via interface, manuals, help", "The designer’s private diary", "Only hardware serial numbers"], 1),
    ("The user’s model is:", ["Always identical to the system image", "How the user understands the system works", "Only the marketing brochure", "The database schema"], 1),
    ("Benefits of conceptualizing early include:", ["Narrow fixation only", "Orientation, open-minded exploration, common ground/terms", "Avoiding questions", "Skipping assumptions"], 1),
    # L5 cognition
    ("Cognition includes:", ["Only memory", "Thinking, remembering, learning, deciding, seeing, reading, talking, writing, etc.", "Only typing speed", "Only hardware"], 1),
    ("Kahneman’s “fast vs slow” relates to:", ["Screen refresh rate", "Types of thinking—quick intuitive vs deliberate", "Only animation FPS", "Network ping"], 1),
    ("Understanding cognition helps designers:", ["Ignore errors", "Know what users can/cannot do; explain problems; use theories and methods", "Remove evaluation", "Avoid prototypes"], 1),
    ("Focused vs divided attention refers to:", ["Monitor size", "Whether attention is on one stream or split across tasks", "Only color blindness", "Font licensing"], 1),
    ("Tullis (1987) hotel search example showed:", ["Density always determines speed", "Same information density but grouping/spacing changed search time (5.5s vs 3.2s)", "Color never matters", "Audio always helps"], 1),
    ("Using a handheld phone while driving is presented as:", ["Safe if under speed limit", "Unsafe—attention is diverted; higher crash risk", "Safer than eating", "Recommended for navigation"], 1),
    ("Hands-free phone while driving:", ["Eliminates cognitive load from conversation", "Still risky—conversation competes for attention", "Always safer than passengers", "Is legally banned everywhere"], 1),
    ("Talking to a front passenger can be less dangerous than remote phone because:", ["Passengers talk more", "Passengers can pause when a hazard appears; remote speakers may not", "Phones use less data", "Passengers never speak"], 1),
    ("Perception implications include:", ["Illegible text is fine", "Legible text; distinguishable icons; good contrast", "Only 3D graphics", "No grouping"], 1),
    ("Weller (2004) found grouping with borders:", ["Always slower than color alone", "Could help people locate items faster than weak color contrast grouping in the example", "Irrelevant to search", "Only for audio"], 1),
    ("Too much white space on some pages can:", ["Always help", "Sometimes make finding information harder", "Fix all accessibility issues", "Replace all menus"], 1),
    ("Yellow on black or blue is noted as:", ["Always bad contrast", "Generally acceptable contrast", "Illegal", "Only for print"], 1),
    ("Yellow on green or white is noted as:", ["Ideal pairing", "A poor contrast choice", "Required for accessibility", "Best for long text"], 1),
    ("Recognition memory is generally:", ["Harder than recall for menus", "Easier than pure recall for many interface tasks", "Impossible with icons", "Same as photographic memory always"], 1),
    ("Graphical menus vs command interfaces:", ["Commands need only recognition", "Commands require more recall; graphical options support recognition", "No difference", "Commands are always faster for novices"], 1),
    ("Miller’s 7±2 is cautioned because:", ["It is never cited", "Users scanning lists/tabs may not need to hold all items in working memory", "It bans more than 3 tabs", "It applies only to audio"], 1),
    ("Password managers and biometrics are discussed as ways to:", ["Increase memory load", "Reduce user memory load for authentication", "Remove security", "Ban MFA"], 1),
    ("MFA (multi-factor authentication) increases steps because:", ["Designers hate users", "Increased security concerns", "It is always faster", "It removes passwords"], 1),
    ("Incidental learning is:", ["Studying for exams only", "Learning without trying—e.g., recognizing faces", "Always harder than intentional", "Impossible with technology"], 1),
    ("People often prefer to learn by:", ["Reading manuals only", "Doing rather than only following manuals", "Avoiding practice", "Memorizing error codes"], 1),
    ("Listening vs reading (slides):", ["Reading always uses more cognitive effort than listening", "Listening can require less cognitive effort than reading/speaking; reading can be faster than speaking/listening", "They are identical always", "Listening is impossible"], 1),
    ("Voice UIs mentioned include:", ["Only Fortran compilers", "Google Voice, Siri, Alexa", "Only punch cards", "Only SQL"], 1),
    ("Design for synthetic speech voices:", ["Use very long sentences", "Keep speech-based menus/instructions short; intonation matters", "Never show text", "Match human voice exactly always"], 1),
    # L6 social
    ("Social technologies:", ["Stop all social behavior", "Extend social interaction when people are apart—effects vary", "Replace all face-to-face permanently with no issues", "Only work offline"], 1),
    ("Sacks Rule 1: next speaker can be chosen when:", ["Silence only", "Current speaker asks opinion, question, or request", "Always random", "Never"], 1),
    ("Sacks Rule 2:", ["Speaker must never stop", "Another person self-selects to speak", "Only robots speak", "Music plays"], 1),
    ("Sacks Rule 3:", ["Always interrupt", "Current speaker may continue", "No turn-taking", "Only text"], 1),
    ("Back-channel cues (uh-huh, mm) signal:", ["End conversation", "Continue / listening", "Anger only", "Logout"], 1),
    ("VideoWindow (Bellcore, 1989) was:", ["A text-only BBS", "A large shared video/audio “window” between distant sites", "A spreadsheet", "A robot arm"], 1),
    ("Findings from VideoWindow use included:", ["Users ignored the tech completely", "People talked a lot about the system; often spoke more to co-located others; moving “closer” on video could break AV pickup", "Perfect replacement for travel", "No social effects"], 1),
    ("Telepresence robots can help:", ["Only gaming", "Remote attendance at conferences, schools, museums", "Cooking food", "Replacing all law"], 1),
    ("Co-presence coordination uses:", ["Only email", "Verbal/nonverbal communication, schedules/rules, shared representations", "Random actions", "No awareness"], 1),
    ("Peripheral awareness is:", ["Only VR resolution", "Tracking what happens at the edge of attention—overhearing/overseeing", "Ignoring teammates", "Only CPU load"], 1),
    ("Situational awareness is important for:", ["Only icon fonts", "Understanding surroundings to act appropriately—e.g., ATC, operating theatre", "Choosing wallpaper", "Picking emojis"], 1),
    ("The ReflectTable used LEDs to show:", ["Room temperature", "How much each group member spoke", "Wi-Fi password", "Battery levels"], 1),
    ("Social engagement often involves:", ["Only paid labor", "Voluntary participation and exchange—e.g., retweets connecting many people", "No online hubs", "Banning collaboration"], 1),
    ("The Eugenia Kuyda chatbot example raises:", ["Only performance tuning", "Ethics of simulating a deceased person via their texts", "Keyboard layout", "Database indexing"], 1),
    # L7 emotion
    ("Visceral design concerns:", ["Meaning and culture after months", "Immediate look, feel, sound", "Only server logs", "Only APIs"], 1),
    ("Behavioral design concerns:", ["Only brand slogans", "Use and usability", "Only packaging in stores", "Only stock price"], 1),
    ("Reflective design concerns:", ["Only button pixels", "Meaning and personal value of a product", "Only RAM size", "Only contrast ratios"], 1),
    ("Expressive interfaces can:", ["Never annoy users", "Be fun and reassuring but sometimes intrusive or annoying", "Remove all color", "Ban animation"], 1),
    # L9 data gathering
    ("Triangulation means:", ["One method only", "Looking at data from more than one perspective or type", "Only surveys", "Skipping pilot studies"], 1),
    ("Pilot studies are:", ["Full production launch", "Small trial runs before the main study", "Illegal", "Only for quantitative data"], 1),
    ("Unstructured interviews are:", ["Fully scripted", "Not directed by a script—rich but less replicable", "Only online", "Only closed questions"], 1),
    ("Structured interviews are:", ["Never repeatable", "Tightly scripted—replicable but may lack richness", "Always focus groups", "Only observation"], 1),
    ("Semi-structured interviews:", ["No plan at all", "Guided by script but allow deeper follow-up—balance richness and replicability", "Only yes/no", "Cannot record"], 1),
    ("Focus groups are:", ["Solo diary studies", "Group interviews, often 3–10 people, facilitated", "Only A/B tests", "Only automated logs"], 1),
    ("Interview questions should avoid:", ["Clear wording", "Leading questions, jargon, long compound sentences, unconscious bias", "Open questions", "Consent forms"], 1),
    # Misc cross-lecture
    ("The course code and topic of Lecture 1 slides:", ["CP101 Programming", "CP650 exploring multidisciplinary nature of interaction design", "MBA finance", "Physics 12"], 1),
    ("Double Diamond is mentioned as:", ["A keyboard layout", "A design process visualization (discover/define/develop/deliver style)", "A color palette", "A legal test"], 1),
    ("Stakeholders help designers:", ["Ignore users", "Identify groups to include; users may not know what is technically possible", "Remove requirements", "Skip problem space"], 1),
    ("Choosing among alternatives can use:", ["Only coin flip", "Technical feasibility, user evaluation, prototypes, A/B testing, quality thresholds", "Only aesthetics", "No metrics"], 1),
    ("Open-source resources for ID can include:", ["Only paid Adobe only", "Pattern libraries, frameworks like Bootstrap, community components", "No code", "Only proprietary OS"], 1),
    ("Documentation in Agile UX should be:", ["Always maximal", "Minimal but sufficient—question audience and duplication", "Forbidden", "Only paper"], 1),
    ("Many smartphones apps had how many distinct user types in one cited study?", ["About 10", "382", "2", "Unlimited known exactly"], 1),
    ("When creating interfaces, should you always show at most 7 menu items because of Miller?", ["Yes, always", "Not necessarily—scanning changes the demand on memory", "Only on mobile", "Only for experts"], 1),
    ("Remote conversation media (email, IM, video) compared to F2F:", ["Have identical breakdowns always", "May follow some similar rules but have different breakdowns and repairs", "Never need repair", "Ban turn-taking"], 1),
    ("Anthropomorphism in Lecture 7 overview is:", ["Never discussed", "Listed as a topic with pros and cons", "Only about fonts", "Only about SQL"], 1),
    ("Persuasive technologies aim to:", ["Only display errors", "Influence behavior and habits", "Remove all feedback", "Ban mobile"], 1),
    ("Affective computing relates to:", ["Only file systems", "Detecting/responding to emotion; improving with ML", "Only RAM timing", "Printer drivers"], 1),
    ("Ortony et al. model in slides connects emotion to:", ["Only hue saturation", "How emotional state changes thinking and focus", "Only SQL joins", "CPU scheduling"], 1),
    ("Automatic vs conscious emotions differ in:", ["They are identical", "Speed and duration—automatic rapid; conscious slow/long-lasting", "Only hardware", "Only licensing"], 1),
    ("High cognitive load designs should:", ["Add long complicated procedures", "Avoid unnecessarily long/complicated procedures", "Hide all labels", "Use only recall"], 1),
    ("Design for memory should favor:", ["Arbitrary codes only", "Recognition over recall where possible; folders/categories/colors/timestamps", "More secret questions always", "No search"], 1),
    ("Multitasking in hospital settings is used to illustrate:", ["That multitasking never happens", "Need for new attention/scanning strategies with multiple live data sources", "That screens are illegal", "That paper is banned"], 1),
    ("Invisible automatic controls can:", ["Always improve usability", "Make interfaces harder when users cannot see what will happen", "Remove need for feedback", "Guarantee safety"], 1),
    ("Dyslexia is noted as affecting:", ["Only hearing", "Difficulties with written word recognition/understanding", "Only motor skills", "Only color vision"], 1),
    ("Natural-language / chatbot systems let users:", ["Only draw pixels", "Type questions and receive text responses", "Only use punch cards", "Never get help"], 1),
    ("Social media questions in Lecture 6 include:", ["Only GPU specs", "Overlap of online vs real-life friends; changing etiquette", "Only typography", "Compiler choice"], 1),
    ("Turkle (2015) is cited regarding:", ["Database indexes", "Negative impacts on everyday conversation", "CPU pipelines", "Keyboard matrices"], 1),
    ("Collaborative tools like Miro support:", ["Only solo work", "Shared visual collaboration", "Only batch jobs", "Hardware soldering"], 1),
    ("Co-presence technologies can include:", ["Only email", "Tabletops, whiteboards, shared spaces", "Only fax", "Only spreadsheets"], 1),
    ("“Being cool about disability” slide stresses prosthetics can be:", ["Only ugly and hidden", "Desirable and fashionable, changing language (“wearing their wheels”)", "Illegal", "Only metal"], 1),
    ("Activity-centered design is:", ["Same as ignoring tasks", "One alternative approach named alongside user-centered and systems design", "Only for games", "Banned in Agile"], 1),
    ("Lean UX is given as an example of:", ["Removing all research", "Ongoing user research program over time", "Only waterfall", "Only hardware"], 1),
    ("Crowdsourcing design ideas is listed under:", ["Security audits", "Degrees of user involvement", "Compiler design", "Payment only"], 1),
    ("Citizen science appears under:", ["Only corporate tax", "Online/user contribution models", "GPU drivers", "Font licensing"], 1),
    ("Interface styles listed include (among others):", ["Only CLI", "Command, speech, form fill-in, graphical, web, gesture, AR, etc.", "Only punch cards", "Only Morse"], 1),
    ("Hybrid conceptual models can:", ["Support only one action forever", "Support multiple ways to carry out the same actions", "Remove consistency", "Ban voice"], 1),
    ("Frameworks (Lecture 4) provide:", ["Only source code", "Interrelated concepts/questions—e.g., DiCoT for team tech use", "Only RGB values", "Only budgets"], 1),
    ("Models simplify phenomena to:", ["Remove all prediction", "Help predict and evaluate designs", "Replace theory entirely", "Ban users"], 1),
    ("Visions (Lecture 4) serve as:", ["Bug trackers", "Imagined futures framing R&D—smart cities, human-centered AI, etc.", "Only payroll", "Only fonts"], 1),
]

# Add more numeric/detail MCQs
extra = [
    ("How many usability goals are explicitly listed? (effective, efficient, safe, utility, learn, remember)", ["Four", "Five", "Six", "Seven"], 2),
    ("Lecture 1 lists how many “characteristics of interaction design” bullets?", ["Two", "Three", "Four", "Five"], 1),
    ("Norman’s emotional design levels count:", ["Two", "Three", "Four", "Five"], 1),
    ("Components of Norman’s conceptual model communication triangle:", ["Two", "Three", "Four", "Five"], 1),
]
MCQS.extend(extra)

# Write questions
for i, (q, opts, correct) in enumerate(MCQS, 1):
    p = doc.add_paragraph()
    p.add_run(f"Q{i}. ").bold = True
    p.add_run(q)
    labels = ["A", "B", "C", "D"]
    for j, opt in enumerate(opts):
        doc.add_paragraph(f"{labels[j]}) {opt}", style="List Bullet")

doc.add_page_break()
doc.add_heading("Answer Key", level=1)
for i, (_, opts, correct) in enumerate(MCQS, 1):
    letter = ["A", "B", "C", "D"][correct]
    doc.add_paragraph(f"Q{i}: {letter}", style="List Number")

out_path = "/Users/umeraamir/Downloads/UIUX/CP650_Study_Guide_and_MCQs.docx"
doc.save(out_path)
print(f"Saved {out_path} with {len(MCQS)} questions.")
